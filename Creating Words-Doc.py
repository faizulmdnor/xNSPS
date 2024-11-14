from docx import Document

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Installation Guide for Python, PyCharm, and SQL Server Express in Windows Environment', 0)

# Section 1: Install Python
doc.add_heading('1. Install Python', level=1)
doc.add_paragraph("1. Download Python:\n"
    "   - Go to the official Python website: https://www.python.org/downloads/.\n"
    "   - Click Download Python 3.x.x (the latest stable version)."
    "\n\n2. Install Python:\n"
    "   - Open the downloaded installer.\n"
    "   - Important: Check the box labeled Add Python to PATH at the bottom of the installation window.\n"
    "   - Click Install Now to start the installation.\n"
    "   - After installation, open Command Prompt and type python --version to confirm Python was installed and added to PATH."
)

# Section 2: Install SQL Server Express
doc.add_heading('2. Install SQL Server Express', level=1)
doc.add_paragraph("1. Download SQL Server Express:\n"
    "   - Go to the official SQL Server Express download page: https://www.microsoft.com/en-us/sql-server/sql-server-downloads.\n"
    "   - Select SQL Server Express and click Download now."
    "\n\n2. Install SQL Server Express:\n"
    "   - Open the downloaded installer.\n"
    "   - Choose Basic Installation for a standard setup.\n"
    "   - Follow the on-screen instructions to complete the installation.\n"
    "   - Take note of the SQL Server instance name provided during installation (e.g., SQLEXPRESS)."
    "\n\n3. Configure SQL Server Express:\n"
    "   - Launch SQL Server Management Studio (SSMS) to connect to the SQL Server instance.\n"
    "   - In the Connect to Server window, select Server type as 'Database Engine' and Server name as 'Your-PC-Name\\SQLEXPRESS'.\n"
    "   - Authentication: Choose Windows Authentication.\n"
    "   - Click Connect to ensure that your SQL Server instance is running."
)

# Section 3: Install PyCharm
doc.add_heading('3. Install PyCharm', level=1)
doc.add_paragraph("1. Download PyCharm:\n"
    "   - Go to the PyCharm website: https://www.jetbrains.com/pycharm/download/.\n"
    "   - Download PyCharm Community Edition (free) or Professional Edition if you have a license."
    "\n\n2. Install PyCharm:\n"
    "   - Open the downloaded installer.\n"
    "   - Select the installation folder and follow the on-screen instructions.\n"
    "   - In the installation options, you may want to:\n"
    "       - Check Create Desktop Shortcut.\n"
    "       - Check Add to PATH if prompted.\n"
    "       - Associate .py files with PyCharm for easier file management.\n"
    "   - Click Install to complete the installation."
)

# Section 4: Configure a Virtual Environment (.venv) in PyCharm
doc.add_heading('4. Configure a Virtual Environment (.venv) in PyCharm', level=1)
doc.add_paragraph("1. Open PyCharm and Create/Open a Project:\n"
    "   - Open PyCharm.\n"
    "   - Select New Project or Open an existing project."
    "\n\n2. Configure a Virtual Environment:\n"
    "   - In the New Project window, select Location for your project.\n"
    "   - Under Python Interpreter, select New environment using and choose Virtualenv.\n"
    "   - PyCharm will automatically detect the installed Python version, but you can specify a different version if needed.\n"
    "   - Location for .venv: PyCharm will create a .venv folder within your project directory by default.\n"
    "   - Click Create to set up the virtual environment."
    "\n\n3. Activate .venv:\n"
    "   - The .venv environment should now be automatically activated for your project.\n"
    "   - To verify, open Terminal in PyCharm (bottom left) and type:\n"
    "       ```bash\n"
    "       python --version\n"
    "       ```\n"
    "     You should see the Python version in the .venv environment."
    "\n\n4. Install Packages in .venv (Optional):\n"
    "   - You can install packages directly within the .venv environment by opening the PyCharm Terminal and using pip, like so:\n"
    "       ```bash\n"
    "       pip install <package-name>\n"
    "       ```"
    "\n\n5. Managing Virtual Environment Settings:\n"
    "   - In PyCharm, you can go to File > Settings > Project: [Project Name] > Python Interpreter to view and manage packages installed in .venv."
)

# Summary
doc.add_heading('Summary', level=1)
doc.add_paragraph("With this setup, you now have:\n"
    "1. Python installed and ready for development.\n"
    "2. SQL Server Express installed and configured.\n"
    "3. PyCharm set up with a virtual environment (.venv) for managing dependencies and ensuring consistent project setups."
)

# Save document
doc_path = "../xNSPS/Documents/Installation_Guide_Python_PyCharm_SQLExpress.docx"
doc.save(doc_path)

doc_path
